import os
import io
from uuid import uuid4
from datetime import datetime

import streamlit as st
from openai import OpenAI
from langdetect import detect

import pdfplumber
from docx import Document

# ============ Streamlit 基础配置（必须是第一个 st 调用） ============
st.set_page_config(
    page_title="AI 智能简历优化",
    page_icon="🧠",
    layout="wide",
)

# ============ 安全导入 analytics（Google Sheet & Slack） ============
try:
    from analytics import log_event, log_feedback, log_error
except Exception:
    # 如果 analytics 还没配置好，不阻塞正常功能
    def log_event(*args, **kwargs):
        pass

    def log_feedback(*args, **kwargs):
        pass

    def log_error(*args, **kwargs):
        pass


# ============ OpenAI 客户端 ============
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
MODEL_NAME = os.getenv("MODEL_NAME", "gpt-4o-mini")

if OPENAI_API_KEY:
    client = OpenAI(api_key=OPENAI_API_KEY)
else:
    client = None

# ============ Session 级别信息 ============
if "sid" not in st.session_state:
    st.session_state["sid"] = str(uuid4())

SESSION_ID = st.session_state["sid"]

# 首次打开页面埋点
log_event(
    "page_view",
    {
        "sid": SESSION_ID,
        "ts": datetime.utcnow().isoformat(),
        "page": "resume_optimizer",
    },
)


# ============ 一些工具函数 ============

def read_docx(file_bytes: bytes) -> str:
    bio = io.BytesIO(file_bytes)
    doc = Document(bio)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


def read_pdf(file_bytes: bytes) -> str:
    text_chunks = []
    bio = io.BytesIO(file_bytes)
    with pdfplumber.open(bio) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            t = t.strip()
            if t:
                text_chunks.append(t)
    return "\n\n".join(text_chunks)


def detect_lang(text: str) -> str:
    try:
        lang = detect(text)
        if lang.startswith("zh"):
            return "zh"
        else:
            return "en"
    except Exception:
        # 默认中文
        return "zh"


def build_prompt(
    cv_text: str,
    jd_text: str,
    focus_tags,
    custom_points: str,
    need_cover_letter: bool,
) -> tuple[str, str]:
    """
    返回 (system_prompt, user_prompt)
    """
    lang = detect_lang(cv_text + "\n" + jd_text)

    if lang == "zh":
        system_prompt = (
            "你是一名资深求职顾问，擅长根据候选人的简历与目标职位 JD，"
            "提升简历匹配度与专业度，同时在需要时撰写高质量求职信。"
            "你需要在保留事实真实性的前提下，优化表述、量化成果、突出与 JD 的匹配度。"
        )
        focus_text = "、".join(focus_tags) if focus_tags else "综合优化"
        user_prompt = f"""
【任务语言】请全程使用与候选人简历相同的语言（当前自动识别为：{"中文" if lang=="zh" else "英文"}）。

【优化重点】{focus_text}

【候选人原始简历】
{cv_text}

【目标职位 JD 或 特别优化指令】
{jd_text}

【自定义增强点（如果为空可以忽略）】
{custom_points or "（无）"}

【输出要求】
1. 先输出【优化后简历】，按照常见简历结构分段：
   - 个人信息（不要虚构联系方式）
   - 教育背景
   - 实习 / 工作经历（每段经历用要点列出，突出职责 + 量化成果 + 使用技能）
   - 项目经历（如有）
   - 技能 & 证书
2. 请特别注意：
   - 不要虚构并不存在的公司、学校、证书或日期；
   - 可以对已有经历进行更专业的表达和重组；
   - 尽量保留原本的关键信息，但避免啰嗦。
3. 如果用户勾选了生成求职信，请在最后再输出一个【求职信】模块：
   - 用 1~1.5 页左右篇幅；
   - 说明候选人与该职位的匹配度、代表性经历和动机。
4. 输出格式用清晰的小标题和项目符号，适合直接复制到 Word 中使用。
"""
    else:
        system_prompt = (
            "You are an experienced career consultant. "
            "Given a candidate's CV and a target job description, "
            "you will rewrite and improve the CV to better match the role, "
            "while keeping all information truthful. "
            "Optionally, you will also draft a tailored cover letter."
        )
        focus_text = ", ".join(focus_tags) if focus_tags else "overall optimisation"
        user_prompt = f"""
[Language] Please respond in the same language as the candidate's CV (currently detected as: {"Chinese" if lang=="zh" else "English"}).

[Focus]
{focus_text}

[Original CV]
{cv_text}

[Target Job Description or Extra Instructions]
{jd_text}

[Custom Emphasis / Extra Points]
{custom_points or "(none)"}

[Output Requirements]
1. First output an **Improved CV**:
   - Use standard sections (Profile, Education, Experience, Projects, Skills, Certifications, etc.).
   - For each experience, use bullet points focusing on responsibilities + quantified impact + skills/tech stack.
   - Do NOT fabricate employers, schools, degrees, or dates.
2. You may rephrase and reorganise content for clarity and impact, but do not invent fake achievements.
3. If the user has requested a cover letter, then add a **Cover Letter** section at the end:
   - About 1 page.
   - Clearly link the candidate's experience to the role requirements.
4. Make the structure easy to copy-paste into Word.
"""

    if not need_cover_letter:
        # 提醒模型可以忽略求职信部分
        if lang == "zh":
            user_prompt += "\n（本次用户没有勾选生成求职信，如无特别需要可省略【求职信】模块。）"
        else:
            user_prompt += "\n(The user did NOT request a cover letter this time; you may omit the Cover Letter section.)"

    return system_prompt, user_prompt


def call_openai(cv_text: str, jd_text: str, focus_tags, custom_points: str, need_cover_letter: bool) -> str:
    if not client:
        raise RuntimeError("OpenAI client not initialised. Please check OPENAI_API_KEY in Secrets.")

    system_prompt, user_prompt = build_prompt(
        cv_text=cv_text,
        jd_text=jd_text,
        focus_tags=focus_tags,
        custom_points=custom_points,
        need_cover_letter=need_cover_letter,
    )

    resp = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.4,
    )
    return resp.choices[0].message.content.strip()


def make_docx(text: str) -> bytes:
    """将纯文本写入一个简单的 docx，返回二进制内容。"""
    doc = Document()
    for block in text.split("\n\n"):
        p = doc.add_paragraph()
        for line in block.split("\n"):
            p.add_run(line)
        # 额外空行交给 split 处理
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


# ============ UI 布局 ============

# 左侧侧边栏
with st.sidebar:
    st.markdown("### 设置")

    st.caption("（左侧选项仅影响生成时的强调方向）")

    focus_options = ["业务影响", "沟通协作", "技术深度", "数据分析", "领导力潜力"]
    focus_tags = st.multiselect("精修侧重（可多选）", focus_options, default=["业务影响"])

    custom_points = st.text_area(
        "增强点（可自定义）",
        value="例如：强调数据分析/量化成果；突出与目标岗位的匹配；或写作风格要求等…",
        height=120,
    )

    need_cover_letter = st.checkbox("生成求职信（Cover Letter）", value=True)

    enable_ocr = st.checkbox("启用 OCR（扫描 PDF）", value=False)

    st.caption("仅供个人求职使用，禁止商用与爬取。")

# 右侧主区域
st.markdown("## 🧠 AI 智能简历优化")

st.markdown(
    "上传简历，AI 将根据 JD 一键优化；可选生成求职信（Cover Letter，语言自动随简历）。"
)

col_left, col_right = st.columns([1.05, 1.0])

with col_left:
    st.markdown("#### 上传简历（PDF 或 DOCX）")
    uploaded_file = st.file_uploader(
        "Drag and drop file here",
        type=["pdf", "docx"],
        label_visibility="collapsed",
    )
    st.caption("支持 PDF / DOCX，单文件 ≤ 50MB；扫描件可启用 OCR。")

with col_right:
    st.markdown("#### 粘贴目标职位 JD 或 优化指令（可批量、用分隔）")
    jd_text = st.text_area(
        "例如：Actuarial graduate role at Deloitte. 也可以直接写优化指令（如强调哪些技能、写作风格等）",
        height=200,
        label_visibility="collapsed",
    )

st.markdown("---")
st.info("💡 提示：可在左侧设置“精修侧重/增强点”；若 PDF 为扫描件，可开启 OCR。")

generate_btn = st.button("🚀 一键生成", use_container_width=True, type="primary")

st.markdown("---")

# 反馈入口
with st.expander("💬 提交反馈 / 功能建议（可选）"):
    fb_col1, fb_col2 = st.columns([2, 1])
    with fb_col1:
        feedback_text = st.text_area("反馈内容（例如：哪里好用 / 哪里有 bug / 希望新增什么功能）", height=120)
        contact = st.text_input("联系方式（可选，方便我回复你，例如邮箱/微信）")
    with fb_col2:
        if st.button("提交反馈"):
            if feedback_text.strip():
                log_feedback(
                    {
                        "sid": SESSION_ID,
                        "ts": datetime.utcnow().isoformat(),
                        "feedback": feedback_text.strip(),
                        "contact": contact.strip(),
                    }
                )
                st.success("感谢反馈！我会尽快查看并优化。")
            else:
                st.warning("请输入一些反馈内容再提交～")


# 页面底部版权
st.caption("© 2025 AI Resume Optimizer | 仅供个人求职使用，禁止商用与爬取。")


# ============ 主逻辑：点击 “一键生成” ============
def handle_generate():
    if not uploaded_file:
        st.warning("请先上传一份 PDF 或 DOCX 简历。")
        return

    # 文件大小限制 50MB
    if uploaded_file.size > 50 * 1024 * 1024:
        st.error("文件过大，请控制在 50MB 以内。")
        return

    if not jd_text.strip():
        st.warning("建议粘贴至少一个目标职位 JD 或 优化指令，这样效果会更好哦。")

    # 读取文件
    file_bytes = uploaded_file.read()
    file_name = uploaded_file.name.lower()

    try:
        if file_name.endswith(".docx"):
            cv_text = read_docx(file_bytes)
        elif file_name.endswith(".pdf"):
            if enable_ocr:
                st.info("当前版本暂未集成 OCR 引擎，将先尝试直接识别 PDF 文本。")
            cv_text = read_pdf(file_bytes)
        else:
            st.error("当前仅支持 PDF 和 DOCX 格式。")
            return
    except Exception as e:
        log_error(
            "file_parse_error",
            {
                "sid": SESSION_ID,
                "file_name": uploaded_file.name,
                "error": str(e),
            },
        )
        st.error("读取简历文件时出错，请确认文件是否正常或稍后重试。")
        return

    if not cv_text.strip():
        st.error("没有从简历中读取到有效文本，可能是扫描件或加密文件。")
        return

    with st.spinner("AI 正在为你优化简历，请稍候…"):
        try:
            result_text = call_openai(
                cv_text=cv_text,
                jd_text=jd_text,
                focus_tags=focus_tags,
                custom_points=custom_points,
                need_cover_letter=need_cover_letter,
            )
        except Exception as e:
            log_error(
                "openai_error",
                {
                    "sid": SESSION_ID,
                    "error": str(e),
                },
            )
            st.error("调用 AI 接口时出错，请检查 API Key 或稍后重试。")
            return

    # 成功埋点
    log_event(
        "generate_success",
        {
            "sid": SESSION_ID,
            "ts": datetime.utcnow().isoformat(),
            "file_name": uploaded_file.name,
            "file_size": uploaded_file.size,
            "need_cover_letter": need_cover_letter,
            "focus_tags": focus_tags,
        },
    )

    st.markdown("### ✅ 生成结果（可直接复制或下载为 Word）")
    st.markdown(result_text)

    # 导出 Word
    docx_bytes = make_docx(result_text)
    safe_name = os.path.splitext(uploaded_file.name)[0]
    export_filename = f"{safe_name}_AI优化版.docx"

    st.download_button(
        label="⬇️ 下载 Word 版本（DOCX）",
        data=docx_bytes,
        file_name=export_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if generate_btn:
    handle_generate()